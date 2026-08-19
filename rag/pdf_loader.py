from pypdf import PdfReader
from docx import Document as DocxDocument

from langchain_core.documents import Document


def load_pdf(uploaded_file):

    documents = []

    reader = PdfReader(uploaded_file)

    for page_number, page in enumerate(reader.pages):

        text = page.extract_text()

        if text and text.strip():

            document = Document(
                page_content=text,
                metadata={
                    "source": uploaded_file.name,
                    "page": page_number + 1,
                    "file_type": "pdf"
                }
            )

            documents.append(document)

    return documents


def load_docx(uploaded_file):

    documents = []

    doc = DocxDocument(uploaded_file)

    paragraphs = []

    for paragraph in doc.paragraphs:

        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    full_text = "\n".join(paragraphs)

    if full_text:

        document = Document(
            page_content=full_text,
            metadata={
                "source": uploaded_file.name,
                "page": 1,
                "file_type": "docx"
            }
        )

        documents.append(document)

    return documents


def load_resumes(uploaded_files):

    documents = []

    for uploaded_file in uploaded_files:

        file_name = uploaded_file.name.lower()

        if file_name.endswith(".pdf"):

            documents.extend(
                load_pdf(uploaded_file)
            )

        elif file_name.endswith(".docx"):

            documents.extend(
                load_docx(uploaded_file)
            )

        else:

            continue

    return documents


def load_single_pdf(uploaded_file):

    if uploaded_file is None:
        return ""

    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):

        documents = load_pdf(uploaded_file)

    elif file_name.endswith(".docx"):

        documents = load_docx(uploaded_file)

    else:

        return ""

    return "\n\n".join(
        document.page_content
        for document in documents
    )